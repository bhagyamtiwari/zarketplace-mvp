#!/usr/bin/env python3
"""Generates editable Word versions of both seller documents.

Same copy as the PDFs, from scripts/seller_content.py, so the two formats cannot
say different things. This is the editable master: open in Google Docs, swap the
image placeholders for real photos, change wording.

It deliberately does not chase the PDF's look. Word cannot hold letterspaced
Inter at 6pt across a reflowing page, and pretending otherwise produces a
document that breaks the moment you drop a photo in. What it keeps is the
structure, the hierarchy, and black-and-white discipline.

Run:  python3 scripts/seller_docx.py
Out:  public/zarketplace-welcome-pack.docx
      public/zarketplace-seller-guide.docx
"""
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import seller_content as T  # noqa: E402

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, hexcolour):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolour)
    cell._tc.get_or_add_tcPr().append(el)


def run(par, text, size=10, bold=False, colour=BLACK, caps=False, spacing=None):
    r = par.add_run(text.upper() if caps else text)
    r.font.name = "Inter"
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = colour
    # Inter is unlikely to be installed for the reader; Arial is the fallback
    # that keeps the grotesque feel rather than dropping to a serif.
    rpr = r._element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), "Inter")
    rpr.append(rf)
    if spacing:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(spacing * 20)))
        rpr.append(sp)
    return r


def para(doc, text="", size=10, bold=False, colour=BLACK, caps=False,
         space_before=0, space_after=6, align=None, spacing=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run(p, text, size, bold, colour, caps, spacing)
    return p


def heading(doc, text):
    """Section header: uppercase, tracked, with a rule under it."""
    p = para(doc, text, size=11, bold=True, caps=True, space_before=16, space_after=2, spacing=1.6)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "BBBBBB")
    bottom.set(qn("w:space"), "4")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def bullet(doc, text, size=9.5, colour=GREY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run(p, text, size, False, colour)
    return p


def slot_box(doc, n, subject, height_in=1.6):
    """A shaded placeholder row: delete it and paste your photo in its place."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade(cell, "F2F2F2")
    cell.height = Inches(height_in)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, f"SLOT {n:02d}  ", 9, True, GREY, spacing=1.4)
    run(p, subject.upper(), 8, False, GREY, spacing=1.0)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, "Replace this box with your photo", 8, False, GREY)
    return t


def dark_block(doc, head, lines, big=None):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade(cell, "000000")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    if big:
        run(p, big + "   ", 26, True, WHITE)
    run(p, head.upper(), 13, True, WHITE, spacing=1.2)
    for ln in lines:
        q = cell.add_paragraph()
        q.paragraph_format.space_after = Pt(2)
        run(q, ln, 9.5, False, RGBColor(0xD0, 0xD0, 0xD0))
    cell.add_paragraph()
    return t


def two_col(doc, left_head, left_items, right_head, right_items):
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for col, head in ((0, left_head), (1, right_head)):
        p = t.rows[0].cells[col].paragraphs[0]
        run(p, head, 9.5, True, BLACK, caps=True, spacing=1.4)
    for col, items in ((0, left_items), (1, right_items)):
        cell = t.rows[1].cells[col]
        cell.paragraphs[0].text = ""
        for i, it in enumerate(items):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run(p, "•  ", 9.5, True, BLACK)
            run(p, it, 9.5, False, GREY)
    return t


def steps_list(doc, items):
    for i, (label, sub) in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run(p, f"{i}.  ", 11, True, BLACK)
        run(p, label, 11, True, BLACK)
        run(p, f"   {sub}", 9.5, False, GREY)


def setup(doc, title):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.7)
    doc.core_properties.title = title
    doc.core_properties.author = "zarketplace"


# =============================================================================
def build_welcome(path):
    doc = Document()
    setup(doc, T.D1_TITLE)

    para(doc, "zarketplace", size=22, bold=True, space_after=0)
    para(doc, T.D1_KICKER, size=9, bold=True, colour=GREY, caps=True, spacing=2.0, space_after=14)

    slot_box(doc, 1, "Cover image, wide. Rail of clothes or a packed parcel", 1.9)
    para(doc, T.D1_HEAD, size=26, bold=True, space_before=10, space_after=0)
    para(doc, T.D1_WHISPER, size=15, colour=GREY, space_after=6)
    para(doc, T.D1_INTRO, size=10, colour=GREY, space_after=4)

    heading(doc, "Why sell here")
    for _icon, label, line in T.PROMISES:
        p = para(doc, space_after=3)
        run(p, label, 10.5, True, BLACK, caps=True, spacing=1.0)
        run(p, f"   {line}", 9.5, False, GREY)

    heading(doc, "How selling works")
    steps_list(doc, T.HOW_SELLING)

    dark_block(doc, "You keep everything", T.KEEP_LINES, big=T.KEEP_HEAD)

    heading(doc, "One listing = one item")
    two_col(doc, "Allowed", T.ALLOWED, "Not allowed", T.NOT_ALLOWED)
    para(doc, T.ONE_ITEM_NOTE, size=9.5, colour=BLACK, space_before=8)

    doc.add_page_break()

    heading(doc, "Great photos sell faster")
    para(doc, "Six shots of the same garment. Front and back are required.",
         size=9.5, colour=GREY, space_after=8)
    for i, (caption, subject, required) in enumerate(T.PHOTO_SLOTS):
        p = para(doc, space_after=2)
        run(p, f"{caption}   ", 10, True, BLACK, caps=True, spacing=1.0)
        run(p, "REQUIRED" if required else "recommended", 8.5, False, GREY, caps=required)
        slot_box(doc, i + 2, subject, 1.5)
        para(doc, space_after=4)

    two_col(doc, "Use", T.PHOTO_USE, "Skip", T.PHOTO_SKIP)
    para(doc, T.PHOTO_NOTE, size=9.5, colour=BLACK, space_before=8)

    heading(doc, "Good photo, bad photo")
    for i, (title, sub) in enumerate(T.COMPARE):
        p = para(doc, space_after=2)
        run(p, title, 10.5, True, BLACK, caps=True, spacing=1.0)
        run(p, f"   {sub}", 9.5, False, GREY)
        slot_box(doc, 8 + i, title, 1.7)
        para(doc, space_after=4)

    heading(doc, "Describe it accurately")
    para(doc, "  ·  ".join(T.DESCRIBE_CHIPS), size=10.5, bold=True, caps=True, spacing=1.0)
    para(doc, T.DESCRIBE_NOTE, size=9.5, colour=GREY)

    heading(doc, "Price it fairly")
    para(doc, T.PRICE_NOTE, size=9.5, colour=GREY)

    dark_block(doc, T.HELP_HEAD, [T.WHATSAPP] + T.HELP_LINES + [f"{T.EMAIL}   {T.INSTAGRAM}"])

    doc.save(path)
    return path


def build_guide(path):
    doc = Document()
    setup(doc, T.D2_TITLE)

    para(doc, "zarketplace", size=22, bold=True, space_after=0)
    para(doc, T.D2_KICKER, size=9, bold=True, colour=GREY, caps=True, spacing=2.0, space_after=14)

    slot_box(doc, 10, "Cover image, wide. A parcel being handed over", 1.7)
    para(doc, T.D2_HEAD, size=26, bold=True, space_before=10, space_after=0)
    para(doc, T.D2_WHISPER, size=15, colour=GREY, space_after=6)

    heading(doc, "After your item sells")
    steps_list(doc, T.AFTER_SALE)
    para(doc, T.DEADLINE_NOTE, size=9.5, colour=BLACK, space_before=8)

    heading(doc, "Packing")
    for i, name in enumerate(T.PACK_SLOTS):
        para(doc, name, size=10.5, bold=True, caps=True, spacing=1.0, space_after=2)
        slot_box(doc, 11 + i, name, 1.4)
        para(doc, space_after=4)
    for tip in T.PACK_TIPS:
        bullet(doc, tip)

    dark_block(doc, "Payouts", T.PAYOUT_HELD + [""] + T.PAYOUT_THEN)

    doc.add_page_break()

    heading(doc, "Need to change a listing?")
    para(doc, T.EDIT_NOTE, size=9.5, colour=GREY)

    heading(doc, "What isn't allowed")
    for b in T.BANNED:
        bullet(doc, b)
    para(doc, T.BANNED_NOTE, size=9.5, colour=BLACK, space_before=8)

    heading(doc, "Frequently asked")
    for q, a in T.FAQ:
        para(doc, q, size=10.5, bold=True, caps=True, spacing=1.0, space_after=2)
        para(doc, a, size=9.5, colour=GREY, space_after=8)

    heading(doc, "Who pays what")
    t = doc.add_table(rows=len(T.SHIPPING) + 1, cols=2)
    p = t.rows[0].cells[0].paragraphs[0]
    run(p, "Buyer pays shipping at checkout", 9.5, True, BLACK, caps=True, spacing=1.2)
    for i, (label, rate) in enumerate(T.SHIPPING, 1):
        run(t.rows[i].cells[0].paragraphs[0], label, 9.5, False, GREY)
        rp = t.rows[i].cells[1].paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run(rp, rate, 9.5, True, BLACK)
    para(doc, "You pay nothing, unless you switch on free shipping for a listing, in "
              "which case that cost comes out of your payout. Everything else, "
              "including Buyer Protection, is on the buyer.",
         size=9.5, colour=GREY, space_before=8)

    heading(doc, "Before you list")
    for item in T.CHECKLIST:
        p = para(doc, space_after=2)
        run(p, "☐  ", 11, False, BLACK)
        run(p, item, 9.5, False, GREY)

    dark_block(doc, T.CLOSING_HEAD, [T.CLOSING_WHISPER, f"WhatsApp {T.WHATSAPP}"])

    doc.save(path)
    return path


if __name__ == "__main__":
    a = build_welcome(os.path.join(ROOT, "public", "zarketplace-welcome-pack.docx"))
    b = build_guide(os.path.join(ROOT, "public", "zarketplace-seller-guide.docx"))
    print(f"wrote {a}\nwrote {b}")
